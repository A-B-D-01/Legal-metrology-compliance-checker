from pathlib import Path

from docx import Document
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from dotenv import load_dotenv


# --------------------------------------------------
# PATHS
# --------------------------------------------------

PROJECT_ROOT = Path(__file__).resolve().parents[2]

SOURCE_DOC = PROJECT_ROOT / "backend" / "knowledge_base" / "regulatory_rules.docx"

VECTOR_STORE_DIR = (
    PROJECT_ROOT
    / "backend"
    / "vector_stores"
    / "legal_metrology_db"
)


# --------------------------------------------------
# LOAD ENVIRONMENT VARIABLES
# --------------------------------------------------

load_dotenv(PROJECT_ROOT / "backend" / ".env")


# --------------------------------------------------
# READ DOCX
# --------------------------------------------------

def load_docx(path: Path):
    print(f"Reading document: {path}")

    if not path.exists():
        raise FileNotFoundError(
            f"Regulatory document not found:\n{path}"
        )

    doc = Document(str(path))

    paragraphs = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    # Also read tables because regulatory documents
    # sometimes contain important information inside tables.
    for table in doc.tables:
        for row in table.rows:
            row_text = []

            for cell in row.cells:
                text = cell.text.strip()

                if text:
                    row_text.append(text)

            if row_text:
                paragraphs.append(" | ".join(row_text))

    full_text = "\n".join(paragraphs)

    if not full_text.strip():
        raise ValueError(
            "The DOCX file was opened, but no readable text was found."
        )

    print(f"Extracted {len(full_text):,} characters.")

    return full_text


# --------------------------------------------------
# CREATE DOCUMENT CHUNKS
# --------------------------------------------------

def create_chunks(text: str):
    print("Splitting regulatory document into chunks...")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=[
            "\n\n",
            "\n",
            ". ",
            " ",
            "",
        ],
    )

    chunks = splitter.split_text(text)

    documents = []

    for i, chunk in enumerate(chunks):
        documents.append(
            LCDocument(
                page_content=chunk,
                metadata={
                    "source": "regulatory_rules.docx",
                    "chunk_id": i,
                },
            )
        )

    print(f"Created {len(documents)} chunks.")

    return documents


# --------------------------------------------------
# CREATE FAISS VECTOR STORE
# --------------------------------------------------

def create_vector_store(documents):
    print("Creating Google embeddings...")

    embeddings = GoogleGenerativeAIEmbeddings(
        model="gemini-embedding-001"
)

    print("Building FAISS vector store...")

    vector_store = FAISS.from_documents(
        documents,
        embeddings,
    )

    VECTOR_STORE_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    print(f"Saving vector store to:\n{VECTOR_STORE_DIR}")

    vector_store.save_local(
        str(VECTOR_STORE_DIR)
    )

    print("FAISS vector store created successfully.")


# --------------------------------------------------
# MAIN
# --------------------------------------------------

def main():
    print("=" * 60)
    print("LEGAL METROLOGY RAG - DOCUMENT INGESTION")
    print("=" * 60)

    text = load_docx(SOURCE_DOC)

    documents = create_chunks(text)

    create_vector_store(documents)

    print("=" * 60)
    print("INGESTION COMPLETE")
    print("=" * 60)


if __name__ == "__main__":
    main()